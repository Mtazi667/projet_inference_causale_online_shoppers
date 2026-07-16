from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
import sys
from zipfile import ZIP_DEFLATED, ZipFile

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree


# ---------------------------------------------------------------------------
# Paths and design tokens
# ---------------------------------------------------------------------------

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "04_outputs" / "rapport"
FIG_DIR = ROOT / "04_outputs" / "figures"
TABLE_DIR = ROOT / "04_outputs" / "tableaux"
DATA_PATH = ROOT / "02_data" / "raw" / "online_shoppers_intention.csv"
OUTPUT_DOCX = OUT_DIR / "rapport_final_inference_causale_online_shoppers.docx"

# Base preset: narrative_proposal. Named override: academic_serif.
# The override deliberately replaces Calibri/blue accents with a traditional,
# restrained university-report system in Times New Roman and black ink.
FONT_BODY = "Times New Roman"
FONT_MATH = "Cambria Math"
FONT_CODE = "Consolas"
INK = "000000"
MUTED = "666666"
RULE = "B7B7B7"
TABLE_HEADER = "E7E6E6"
TABLE_ALT = "F7F7F7"
WHITE = "FFFFFF"

PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(1)
HEADER_DISTANCE = Inches(0.492)
FOOTER_DISTANCE = Inches(0.492)
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 90, "start": 120, "bottom": 90, "end": 120}


# ---------------------------------------------------------------------------
# Low-level OOXML helpers
# ---------------------------------------------------------------------------

def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(
    run,
    *,
    name: str = FONT_BODY,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = INK,
):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), name)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "fr-CA")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)


def set_style_font(style, name: str, size: float, *, bold=False, italic=False, color=INK):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = rgb(color)
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), name)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "fr-CA")


def set_doc_defaults(doc: Document):
    styles = doc.styles.element
    defaults = styles.find(qn("w:docDefaults"))
    if defaults is None:
        defaults = OxmlElement("w:docDefaults")
        styles.insert(0, defaults)
    r_pr_default = defaults.find(qn("w:rPrDefault"))
    if r_pr_default is None:
        r_pr_default = OxmlElement("w:rPrDefault")
        defaults.append(r_pr_default)
    r_pr = r_pr_default.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        r_pr_default.append(r_pr)
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT_BODY)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "fr-CA")
    sz = r_pr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        r_pr.append(sz)
    sz.set(qn("w:val"), "24")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGINS_DXA.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4" if edge.startswith("inside") else "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), RULE)


def set_table_geometry(table, widths_dxa: list[int]):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table width must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        set_row_cant_split(row)
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_table_borders(table)


def add_numbering_definition(doc: Document, *, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{0xA1100000 + abstract_id:08X}")
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    template = OxmlElement("w:tmpl")
    template.set(qn("w:val"), f"{0xB1100000 + abstract_id:08X}")
    abstract.append(template)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "\uf0b7" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    lvl.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540" if kind == "bullet" else "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540" if kind == "bullet" else "720")
    ind.set(qn("w:hanging"), "280" if kind == "bullet" else "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    numbering_font = "Symbol" if kind == "bullet" else FONT_BODY
    r_fonts.set(qn("w:ascii"), numbering_font)
    r_fonts.set(qn("w:hAnsi"), numbering_font)
    r_fonts.set(qn("w:eastAsia"), numbering_font)
    r_fonts.set(qn("w:cs"), numbering_font)
    r_pr.append(r_fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def add_num_pr(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


def append_field(paragraph, instruction: str, cached_text: str = ""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = cached_text
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def set_page_number_start(section, start: int):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def enable_field_updates(doc: Document):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def harden_ooxml_fonts(path: Path):
    """Remove theme-font ambiguity while preserving explicit Cambria Math runs."""
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    a_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    w = f"{{{w_ns}}}"
    theme_attributes = {
        f"{w}asciiTheme",
        f"{w}hAnsiTheme",
        f"{w}eastAsiaTheme",
        f"{w}cstheme",
    }
    explicit_attributes = [f"{w}ascii", f"{w}hAnsi", f"{w}eastAsia", f"{w}cs"]
    legacy_theme_fonts = {"calibri", "cambria", "aptos", "aptos display", "aptos narrow"}

    with ZipFile(path, "r") as source:
        entries = [(info, source.read(info.filename)) for info in source.infolist()]

    output_entries = []
    for info, data in entries:
        if info.filename.startswith("word/") and info.filename.endswith(".xml"):
            try:
                root = etree.fromstring(data)
            except etree.XMLSyntaxError:
                output_entries.append((info, data))
                continue

            for fonts in root.xpath(".//w:rFonts", namespaces={"w": w_ns}):
                for attr in theme_attributes:
                    fonts.attrib.pop(attr, None)
                explicit = next(
                    (fonts.get(attr) for attr in explicit_attributes if fonts.get(attr)),
                    FONT_BODY,
                )
                if explicit.casefold() in legacy_theme_fonts:
                    explicit = FONT_BODY
                for attr in explicit_attributes:
                    current = fonts.get(attr)
                    if current is None or current.casefold() in legacy_theme_fonts:
                        fonts.set(attr, explicit)

            if info.filename == "word/theme/theme1.xml":
                for xpath in (
                    ".//a:themeElements/a:fontScheme/a:majorFont/a:latin",
                    ".//a:themeElements/a:fontScheme/a:minorFont/a:latin",
                ):
                    for node in root.xpath(xpath, namespaces={"a": a_ns}):
                        node.set("typeface", FONT_BODY)
            data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
        output_entries.append((info, data))

    temporary = path.with_suffix(path.suffix + ".fonttmp")
    with ZipFile(temporary, "w", compression=ZIP_DEFLATED, compresslevel=9) as target:
        for info, data in output_entries:
            target.writestr(info, data)
    temporary.replace(path)


# ---------------------------------------------------------------------------
# Styles and section furniture
# ---------------------------------------------------------------------------

def configure_styles(doc: Document):
    set_doc_defaults(doc)
    styles = doc.styles

    normal = styles["Normal"]
    set_style_font(normal, FONT_BODY, 12, color=INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Heading 1": (15, True, False, 15, 6),
        "Heading 2": (12, True, False, 10, 4),
        "Heading 3": (12, False, True, 8, 3),
    }
    for name, (size, bold, italic, before, after) in heading_specs.items():
        style = styles[name]
        set_style_font(style, FONT_BODY, size, bold=bold, italic=italic, color=INK)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True

    for style_name in ("List Paragraph", "List Bullet", "List Number"):
        style = styles[style_name]
        set_style_font(style, FONT_BODY, 12, color=INK)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.widow_control = True

    caption = styles["Caption"]
    set_style_font(caption, FONT_BODY, 10, italic=True, color=MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(2)
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.keep_together = True

    custom_specs = {
        "Front Matter Heading": (15, True, False, INK),
        "Table Caption": (10, True, False, INK),
        "Source": (9, False, True, MUTED),
        "Equation": (11.2, False, False, INK),
        "Academic Quote": (11.5, False, True, INK),
        "Academic Code": (8.5, False, False, INK),
        "TOC Heading": (15, True, False, INK),
        "TOC 1": (11, False, False, INK),
    }
    for name, (size, bold, italic, color) in custom_specs.items():
        if name in styles:
            style = styles[name]
        else:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        font_name = FONT_MATH if name == "Equation" else FONT_CODE if name == "Academic Code" else FONT_BODY
        set_style_font(style, font_name, size, bold=bold, italic=italic, color=color)

    table_caption = styles["Table Caption"]
    table_caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_caption.paragraph_format.space_before = Pt(6)
    table_caption.paragraph_format.space_after = Pt(3)
    table_caption.paragraph_format.keep_with_next = True
    table_caption.paragraph_format.keep_together = True

    front_heading = styles["Front Matter Heading"]
    front_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    front_heading.paragraph_format.space_before = Pt(15)
    front_heading.paragraph_format.space_after = Pt(6)
    front_heading.paragraph_format.keep_with_next = True
    front_heading.paragraph_format.keep_together = True

    source = styles["Source"]
    source.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    source.paragraph_format.space_before = Pt(3)
    source.paragraph_format.space_after = Pt(8)
    source.paragraph_format.line_spacing = 1.0

    equation = styles["Equation"]
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.space_before = Pt(5)
    equation.paragraph_format.space_after = Pt(7)
    equation.paragraph_format.keep_together = True

    quote = styles["Academic Quote"]
    quote.paragraph_format.left_indent = Inches(0.45)
    quote.paragraph_format.right_indent = Inches(0.45)
    quote.paragraph_format.space_before = Pt(6)
    quote.paragraph_format.space_after = Pt(8)
    quote.paragraph_format.line_spacing = 1.15
    quote.paragraph_format.keep_together = True

    code = styles["Academic Code"]
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(5)
    code.paragraph_format.line_spacing = 1.0

    toc_heading = styles["TOC Heading"]
    toc_heading.paragraph_format.space_before = Pt(0)
    toc_heading.paragraph_format.space_after = Pt(12)
    toc_heading.paragraph_format.keep_with_next = True

    toc1 = styles["TOC 1"]
    toc1.paragraph_format.left_indent = Inches(0)
    toc1.paragraph_format.first_line_indent = Inches(0)
    toc1.paragraph_format.space_after = Pt(3)
    toc1.paragraph_format.line_spacing = 1.0

    table_normal = styles["Normal Table"]
    set_style_font(table_normal, FONT_BODY, 9.3, color=INK)


def configure_section(section, *, front_matter=False):
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.right_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.header_distance = HEADER_DISTANCE
    section.footer_distance = FOOTER_DISTANCE

    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.different_first_page_header_footer = front_matter

    header = section.header
    for p in header.paragraphs:
        p.text = ""
    footer = section.footer
    for p in footer.paragraphs:
        p.text = ""

    if front_matter:
        section.first_page_header.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False
        section.first_page_header.paragraphs[0].text = ""
        section.first_page_footer.paragraphs[0].text = ""
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        append_field(p, "PAGE \\* ROMAN", "ii")
    else:
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_after = Pt(0)
        run = hp.add_run("Inférence causale sur les transactions en ligne")
        set_run_font(run, size=9, italic=True, color=MUTED)
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        append_field(fp, "PAGE", "1")
        set_page_number_start(section, 1)


# ---------------------------------------------------------------------------
# Content helpers
# ---------------------------------------------------------------------------

def add_body(
    doc: Document,
    text: str,
    *,
    keep_with_next: bool = False,
    keep_together: bool = False,
    align=None,
):
    p = doc.add_paragraph(style="Normal")
    if align is not None:
        p.alignment = align
    p.paragraph_format.keep_with_next = keep_with_next
    p.paragraph_format.keep_together = keep_together
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_lead(doc: Document, label: str, text: str):
    p = doc.add_paragraph(style="Normal")
    r = p.add_run(f"{label}. ")
    set_run_font(r, size=12, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=12)
    return p


def add_heading(doc: Document, text: str, level: int, *, page_break=False):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.page_break_before = page_break
    run = p.add_run(text)
    set_run_font(
        run,
        size=15 if level == 1 else 12,
        bold=level in (1, 2),
        italic=level == 3,
    )
    return p


def add_front_heading(doc: Document, text: str):
    p = doc.add_paragraph(style="Front Matter Heading")
    run = p.add_run(text)
    set_run_font(run, size=15, bold=True)
    return p


def add_list_item(doc: Document, text: str, num_id: int, *, label: str | None = None):
    p = doc.add_paragraph(style="List Paragraph")
    add_num_pr(p, num_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_together = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if label:
        r = p.add_run(f"{label} : ")
        set_run_font(r, size=12, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=12)
    return p


def add_equation(doc: Document, text: str, *, size=11.2, after=7):
    p = doc.add_paragraph(style="Equation")
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, name=FONT_MATH, size=size, color=INK)
    return p


def add_quote(doc: Document, text: str):
    p = doc.add_paragraph(style="Academic Quote")
    run = p.add_run(text)
    set_run_font(run, size=11.5, italic=True)
    return p


def add_code_line(doc: Document, text: str):
    p = doc.add_paragraph(style="Academic Code")
    run = p.add_run(text)
    set_run_font(run, name=FONT_CODE, size=8.5)
    return p


def add_table_caption(doc: Document, number: str, title: str):
    p = doc.add_paragraph(style="Table Caption")
    run = p.add_run(f"Tableau {number}. {title}")
    set_run_font(run, size=10, bold=True)
    return p


def add_source(doc: Document, text: str):
    p = doc.add_paragraph(style="Source")
    run = p.add_run(text)
    set_run_font(run, size=9, italic=True, color=MUTED)
    return p
    return p


def add_table(
    doc: Document,
    *,
    number: str,
    title: str,
    headers: list[str],
    rows: list[list[object]],
    widths: list[int],
    numeric_cols: tuple[int, ...] = (),
    font_size: float = 9.3,
    source: str,
    keep_whole: bool = True,
):
    add_table_caption(doc, number, title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = str(value)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = str(value)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            for cell in row.cells:
                set_cell_shading(cell, TABLE_HEADER)
        elif row_index % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, TABLE_ALT)
        for col_index, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                p.alignment = (
                    WD_ALIGN_PARAGRAPH.RIGHT
                    if col_index in numeric_cols
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                p.paragraph_format.keep_together = True
                if row_index == 0 or keep_whole:
                    p.paragraph_format.keep_with_next = True
                for run in p.runs:
                    set_run_font(
                        run,
                        size=font_size,
                        bold=row_index == 0,
                        color=INK,
                    )
    add_source(doc, source)
    return table


def add_figure(
    doc: Document,
    *,
    filename: str,
    number: int,
    title: str,
    source: str,
    alt: str,
    width: float,
    crop_top: int = 0,
    crop_bottom: int = 0,
):
    path = FIG_DIR / filename
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(5)
    p_img.paragraph_format.space_after = Pt(2)
    p_img.paragraph_format.keep_with_next = True
    p_img.paragraph_format.keep_together = True
    run = p_img.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    if crop_top or crop_bottom:
        pic = shape._inline.graphic.graphicData.pic
        blip_fill = pic.blipFill
        src_rect = blip_fill.find(qn("a:srcRect"))
        if src_rect is None:
            src_rect = OxmlElement("a:srcRect")
            blip = blip_fill.find(qn("a:blip"))
            blip.addnext(src_rect)
        if crop_top:
            src_rect.set("t", str(crop_top))
        if crop_bottom:
            src_rect.set("b", str(crop_bottom))
        shape.height = int(
            shape.height * (100000 - crop_top - crop_bottom) / 100000
        )
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", f"Figure {number}. {title}")

    p_cap = doc.add_paragraph(style="Caption")
    p_cap.paragraph_format.keep_with_next = True
    run = p_cap.add_run(f"Figure {number}. {title}")
    set_run_font(run, size=10, italic=True, color=MUTED)
    add_source(doc, source)


def add_toc(doc: Document):
    p = doc.add_paragraph(style="TOC Heading")
    run = p.add_run("Table des matières")
    set_run_font(run, size=15, bold=True)
    toc = doc.add_paragraph()
    toc.paragraph_format.space_after = Pt(0)
    append_field(toc, 'TOC \\o "1-1" \\h \\z \\u', "La table des matières sera mise à jour automatiquement.")


def add_page_break(doc: Document):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def fr(value: float, digits=3) -> str:
    return f"{value:.{digits}f}".replace(".", ",")


# ---------------------------------------------------------------------------
# Report body
# ---------------------------------------------------------------------------

def build_report():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    df = pd.read_csv(DATA_PATH)
    principal = pd.read_csv(TABLE_DIR / "02_estimations_principales.csv")
    robustness = pd.read_csv(TABLE_DIR / "04_analyses_robustesse.csv")
    pag = pd.read_csv(TABLE_DIR / "01_aretes_pag_fci.csv")
    balance = pd.read_csv(TABLE_DIR / "03_equilibre_covariables.csv", index_col=0)

    doc = Document()
    configure_styles(doc)
    enable_field_updates(doc)
    configure_section(doc.sections[0], front_matter=True)

    props = doc.core_properties
    props.title = "Inférence causale sur les transactions en ligne"
    props.subject = "Effet exploratoire d'une session durant le weekend sur la probabilité de transaction"
    props.author = "Projet d'inférence causale"
    props.last_modified_by = "Projet d'inférence causale"
    props.keywords = "inférence causale, Weekend, Revenue, DAG, FCI, AIPW"
    props.comments = ""
    props.created = datetime(2026, 7, 15, tzinfo=timezone.utc)
    props.modified = datetime(2026, 7, 15, tzinfo=timezone.utc)

    # Cover - editorial_cover pattern, adapted to an academic report.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("RAPPORT FINAL")
    set_run_font(r, size=11, bold=True, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Inférence causale sur les\ntransactions en ligne")
    set_run_font(r, size=23, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("Effet exploratoire d'une session durant le weekend\nsur la probabilité de transaction")
    set_run_font(r, size=13, italic=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(72)
    r = p.add_run("Jeu de données : Online Shoppers Purchasing Intention (UCI)")
    set_run_font(r, size=10.5, color=MUTED)

    for text, bold in (
        ("Mohamed Tazi", True),
        ("Cours : Inférence causale", False),
        ("Unité d'analyse : une session de navigation", False),
        ("Juillet 2026", False),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        set_run_font(r, size=10.5, bold=bold, color=INK)

    add_page_break(doc)

    # Abstract and executive reading page.
    add_front_heading(doc, "Résumé")
    add_body(
        doc,
        "Ce rapport étudie une question causale distincte de la prédiction habituelle de Revenue : "
        "parmi les sessions éloignées d'un jour commercial spécial, quel serait l'effet moyen du "
        "fait qu'une session ait lieu durant le weekend plutôt qu'en semaine sur sa probabilité de "
        "se terminer par une transaction\u00A0? L'analyse porte sur 11 079 sessions avec SpecialDay=0. "
        "Elle combine un graphe causal construit à partir de l'ordre temporel, une découverte "
        "exploratoire par FCI, une identification par la porte arrière sous hypothèses et trois "
        "estimateurs ajustés : standardisation, IPW de Hájek et AIPW avec prédictions croisées."
    )
    add_body(
        doc,
        "Le taux observé est de 18,06 % le weekend et de 16,03 % en semaine. L'association brute "
        "est de +2,03 points de pourcentage. L'estimation AIPW principale est de +1,51 point, avec "
        "un intervalle à 95 % de [-0,18 ; 3,19]. Le recouvrement des scores et l'équilibre des "
        "covariables mesurées sont favorables après pondération dans la population restreinte, mais "
        "l'intervalle principal "
        "contient zéro. Les promotions, les prix, les campagnes et l'intention d'achat préalable "
        "ne sont pas observés. La conclusion est donc conditionnelle : les données sont compatibles "
        "avec un faible contraste positif, sans permettre d'établir un effet causal moyen différent "
        "de zéro."
    )
    add_lead(doc, "Mots-clés", "inférence causale; résultats potentiels; DAG; FCI; score de propension; AIPW.")

    bullets_summary = add_numbering_definition(doc, kind="bullet")
    add_list_item(doc, "Population principale : 11 079 sessions avec SpecialDay=0.", bullets_summary)
    add_list_item(doc, "Estimation AIPW sous hypothèses : +1,506 point; IC 95 % [-0,176 ; 3,188].", bullets_summary)
    add_list_item(doc, "SMD absolu maximal : 0,268 avant IPW et 0,040 après IPW.", bullets_summary)
    add_list_item(doc, "Portée : effet par session existante, et non effet sur le trafic total ou le chiffre d'affaires.", bullets_summary)

    add_page_break(doc)
    add_toc(doc)

    main = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(main, front_matter=False)

    # 1. Introduction
    add_heading(doc, "1. Introduction et question causale", 1)
    add_heading(doc, "1.1 Contexte et pertinence", 2)
    add_body(
        doc,
        "Le commerce électronique utilise fréquemment les traces de navigation pour prédire si une "
        "visite se terminera par un achat. Une bonne prédiction ne répond cependant pas à la question "
        "de savoir ce qui changerait si l'on modifiait le moment de la visite. Des variables comme "
        "PageValues ou la durée passée sur certaines pages peuvent être très prédictives tout en "
        "étant impropres à l'ajustement causal, parce qu'elles sont mesurées pendant la session ou "
        "construites à partir d'informations proches de la transaction."
    )
    add_body(
        doc,
        "Le contraste semaine-weekend est pertinent pour une décision appliquée : il peut aider à "
        "formuler une hypothèse sur le calendrier des opérations, des campagnes ou de la capacité du "
        "site. La base reste toutefois observationnelle et anonyme. Elle ne contient ni assignation "
        "aléatoire, ni prix, ni promotion, ni heure exacte. Le rapport sépare donc systématiquement "
        "association observée, identification sous hypothèses et conclusion causale."
    )

    add_heading(doc, "1.2 Question, intervention et portée", 2)
    add_quote(
        doc,
        "Parmi les sessions avec SpecialDay=0, quel serait l'effet moyen du fait qu'une session ait "
        "lieu durant le weekend plutôt qu'en semaine sur la probabilité qu'elle se termine par une "
        "transaction\u00A0?"
    )
    add_body(
        doc,
        "L'unité est une session de navigation. Weekend est une exposition calendaire, plutôt qu'un "
        "traitement réellement administré. L'intervention hypothétique compare deux moments possibles "
        "dans le même mois, l'un en semaine et l'autre durant le weekend, tous deux avec SpecialDay=0, "
        "en conservant le profil pré-session et, dans l'analyse principale, le canal d'acquisition. "
        "Cette définition rend la comparaison "
        "plus précise, mais elle ne décrit pas toutes les façons possibles de faire survenir une "
        "session durant le weekend."
    )
    add_body(
        doc,
        "Le résultat Revenue indique seulement si une transaction a eu lieu. Il ne représente ni un "
        "montant, ni une intention psychologique déclarée. L'effet visé concerne la probabilité "
        "d'achat parmi des sessions déjà présentes. Il ne mesure pas un effet du weekend sur le "
        "nombre de visiteurs, le nombre total de transactions ou le chiffre d'affaires du site."
    )

    # 2. Data and variables
    add_heading(doc, "2. Données, variables et population cible", 1)
    add_heading(doc, "2.1 Provenance et contrôles", 2)
    add_body(
        doc,
        "Le fichier provient du jeu Online Shoppers Purchasing Intention du UCI Machine Learning "
        "Repository, jeu no 468 [1]. Il rassemble des sessions observées sur une année. Selon la "
        "documentation du jeu, chaque session correspond à un utilisateur distinct dans la "
        "construction de la base [1, 2]. La copie locale est protégée par une empreinte SHA-256 "
        "vérifiée avant l'analyse."
    )
    controls = [
        ["Sessions", f"{len(df):,}".replace(",", " ")],
        ["Colonnes", str(df.shape[1])],
        ["Variables explicatives", "17"],
        ["Valeurs manquantes", str(int(df.isna().sum().sum()))],
        ["Lignes exactement identiques", str(int(df.duplicated().sum()))],
        ["Transactions", f"{int(df['Revenue'].sum()):,}".replace(",", " ")],
        ["Taux de transaction", f"{100 * df['Revenue'].mean():.2f} %".replace(".", ",")],
    ]
    add_table(
        doc,
        number="1",
        title="Contrôles de la base brute",
        headers=["Indicateur", "Valeur"],
        rows=controls,
        widths=[6500, 2860],
        numeric_cols=(1,),
        font_size=9.5,
        source="Source : contrôles exécutés dans le notebook final; N = 12 330.",
    )
    add_body(
        doc,
        "Aucune valeur manquante n'est observée. Les 125 lignes exactement identiques sont conservées "
        "dans l'analyse principale, car elles peuvent correspondre à des sessions distinctes ayant les "
        "mêmes valeurs anonymisées. Une analyse de sensibilité les retire afin de vérifier leur influence."
    )
    add_lead(doc, "Empreinte SHA-256", "la copie analysée est identifiée par l'empreinte suivante :")
    add_code_line(doc, "b3055ee355f59134d851d32641183cb4a8b45def7124d2f50442a042f358e0d9")

    add_heading(doc, "2.2 Rôle causal des variables", 2)
    add_body(
        doc,
        "Le choix des variables repose sur leur position supposée dans le mécanisme, et non sur leur "
        "corrélation avec Revenue. Les informations de calendrier et de profil sont disponibles avant "
        "ou au début de la session. Les nombres de pages, les durées, BounceRates et ExitRates sont "
        "mesurés pendant la navigation; Weekend peut donc les influencer. Les inclure dans l'ensemble "
        "d'ajustement risquerait de bloquer une partie de l'effet total ou d'ouvrir un chemin par collider."
    )
    roles = [
        ["Exposition T", "Weekend", "1 = weekend; 0 = semaine"],
        ["Résultat Y", "Revenue", "Transaction binaire; aucun montant monétaire"],
        ["Ajustement principal Z", "Month; VisitorType; Region; OperatingSystems; Browser; TrafficType", "Covariables pré-session supposées"],
        ["Contexte de sélection", "SpecialDay", "Restriction principale à la valeur 0"],
        ["Médiateurs possibles", "Pages consultées; durées; BounceRates; ExitRates", "Variables de navigation exclues de Z"],
        ["Variable proche de Y", "PageValues", "Ordre temporel incertain; exclue conservativement de Z"],
        ["Causes non observées", "Prix; promotions; campagnes; produits; intention préalable", "Menaces à l'échangeabilité"],
        ["Instrument", "Aucun instrument crédible", "Aucune variable ne satisfait les conditions requises"],
    ]
    add_table(
        doc,
        number="2",
        title="Traitement, résultat et rôle des autres variables",
        headers=["Rôle", "Variables", "Justification"],
        rows=roles,
        widths=[1900, 3300, 4160],
        font_size=8.7,
        source="Source : classification causale définie avant l'estimation dans le notebook final.",
    )
    add_body(
        doc,
        "TrafficType demande une attention particulière. Il peut refléter une campagne ou une préférence "
        "préexistante qui influence aussi le moment de la visite, auquel cas il agit comme confondeur ou "
        "proxy; le calendrier peut inversement modifier le canal d'arrivée, auquel cas il devient un "
        "médiateur. Il est inclus dans l'analyse "
        "principale et retiré dans une analyse de sensibilité substantielle."
    )

    add_heading(doc, "2.3 Recouvrement et population cible", 2)
    crosstab = pd.crosstab(df["SpecialDay"], df["Weekend"]).rename(columns={False: "Semaine", True: "Weekend"})
    overlap_rows = [[fr(float(idx), 1), int(row["Semaine"]), int(row["Weekend"])] for idx, row in crosstab.iterrows()]
    add_table(
        doc,
        number="3",
        title="Répartition des sessions selon SpecialDay et Weekend",
        headers=["SpecialDay", "Semaine", "Weekend"],
        rows=overlap_rows,
        widths=[3120, 3120, 3120],
        numeric_cols=(0, 1, 2),
        font_size=9.4,
        source="Source : tabulation de la base brute; N = 12 330.",
    )
    add_body(
        doc,
        "Les valeurs 0,2, 0,4 et 0,6 n'apparaissent qu'en semaine; la valeur 1 n'apparaît que durant "
        "le weekend; la valeur 0,8 ne compte que six sessions de weekend. "
        "Un ajustement sur la base complète forcerait donc des comparaisons sans support dans plusieurs "
        "strates. L'analyse principale retient SpecialDay=0, soit 11 079 sessions : 8 371 en semaine "
        "et 2 708 le weekend. Cette restriction améliore la positivité, mais redéfinit la population "
        "cible aux sessions éloignées d'un jour commercial spécial."
    )
    add_figure(
        doc,
        filename="01_description_et_recouvrement.png",
        number=1,
        title="Taux de transaction observé et recouvrement de SpecialDay",
        source="Source : calculs réalisés dans le notebook final.",
        alt="Deux graphiques : taux de transaction en semaine et durant le weekend, puis effectifs par valeur de SpecialDay.",
        width=6.25,
        crop_top=7500,
    )
    add_body(
        doc,
        "Dans cette population, 489 des 2 708 sessions de weekend aboutissent à une transaction, "
        "contre 1 342 des 8 371 sessions de semaine. Les taux observés sont 18,06 % et 16,03 %, soit "
        "un écart descriptif de +2,03 points. Cet écart ne constitue pas encore un effet causal, car "
        "l'exposition n'est pas randomisée.",
        keep_together=True,
    )

    # 3. Prior knowledge and domain DAG
    add_heading(doc, "3. Connaissances préalables et DAG de domaine", 1)
    add_heading(doc, "3.1 Ordre temporel et mécanismes plausibles", 2)
    add_body(
        doc,
        "Le DAG de domaine est construit avant l'estimation à partir d'un ordre temporel plausible. "
        "Le mois, le profil du visiteur et le canal d'acquisition précèdent la navigation. Weekend "
        "peut modifier le temps disponible, la façon de parcourir le site et la décision d'achat. Les "
        "mesures de navigation peuvent à leur tour influer sur Revenue. PageValues est mesurée à "
        "travers la navigation et construite à partir de valeurs de pages associées aux conversions. "
        "Son ordre précis par rapport à la transaction courante étant incertain, elle est exclue "
        "conservativement de l'ajustement [1, 2]."
    )
    add_body(
        doc,
        "Deux familles de causes latentes sont ajoutées au graphe. Les promotions et les prix peuvent "
        "varier selon le calendrier, modifier le canal d'arrivée et affecter l'achat. L'intention "
        "d'achat préalable peut influencer le moment de la visite, la navigation et la transaction. "
        "Ces nœuds ne sont pas estimés : ils rendent visible la principale hypothèse non vérifiable."
    )

    add_heading(doc, "3.2 Arêtes retenues, interdites et ambiguës", 2)
    constraints = [
        ["Avant la session", "Calendrier/profil → Weekend, trafic ou Revenue", "Le contexte et le profil ne peuvent pas être causés par la transaction courante."],
        ["Pendant la session", "Weekend/TrafficType → Navigation", "Le moment et le canal peuvent modifier le parcours."],
        ["Résultat", "Navigation → Revenue", "Le comportement durant la session peut précéder la transaction."],
        ["Variable proche du résultat", "PageValues liée à la navigation et aux conversions", "Son ordre précis étant incertain, elle n'est pas retenue comme cause pré-session."],
        ["Directions interdites", "Revenue ↛ Month, Weekend, VisitorType, Region, OS, Browser", "Une transaction courante ne remonte pas le temps pour modifier le contexte."],
        ["Relation ambiguë", "Ordre TrafficType–Weekend non résolu", "Le canal peut précéder le jour ou être modifié par le calendrier."],
        ["Confusion latente", "Promotions/prix et intention → Weekend et Revenue", "Ces causes possibles ne sont pas observées."],
    ]
    add_table(
        doc,
        number="4",
        title="Contraintes causales intégrées au raisonnement",
        headers=["Type", "Relation", "Justification"],
        rows=constraints,
        widths=[1800, 3000, 4560],
        font_size=8.7,
        source="Source : connaissances temporelles et fonctionnelles explicitées dans le notebook final.",
    )
    add_body(
        doc,
        "Le graphe présenté à la figure 2 est le graphe de travail final utilisé pour choisir "
        "l'ajustement. Il ne "
        "prétend pas être le vrai mécanisme complet. Sa fonction est de rendre les choix audités : "
        "ajuster sur les causes pré-session mesurées, éviter les variables de navigation et montrer "
        "explicitement les chemins impossibles à bloquer avec la base disponible."
    )
    add_body(
        doc,
        "La position de PageValues y encode son statut aval ou proche du résultat, sans prétendre "
        "établir que la transaction courante cause directement sa valeur.",
        keep_together=True,
    )
    add_figure(
        doc,
        filename="02_dag_de_travail.png",
        number=2,
        title="DAG de domaine utilisé pour raisonner sur l'identification",
        source="Source : construction raisonnée; les flèches pointillées représentent des causes non observées plausibles.",
        alt="DAG reliant calendrier, profil, canal de trafic, weekend, navigation, transaction et PageValues, avec promotions et intention d'achat non observées.",
        width=6.3,
        crop_top=4500,
    )

    add_heading(doc, "3.3 Ensemble d'ajustement proposé", 2)
    add_equation(doc, "Z = {Month, VisitorType, Region, OperatingSystems, Browser, TrafficType}")
    add_body(
        doc,
        "Dans le sous-graphe observé, Z bloque les chemins de porte arrière passant par le calendrier, "
        "le profil et le canal d'acquisition. Aucun instrument crédible ni mécanisme frontdoor complet "
        "n'est disponible. Si les causes latentes dessinées influencent réellement Weekend et Revenue, "
        "Z ne suffit plus : l'effet n'est alors pas garanti comme point-identifié. Le rapport formule "
        "donc les estimations comme causales exploratoires sous hypothèses."
    )

    # 4. Causal discovery
    add_heading(doc, "4. Découverte causale exploratoire", 1)
    add_heading(doc, "4.1 Pourquoi FCI\u00A0?", 2)
    add_body(
        doc,
        "Le DAG n'est que partiellement connu et des causes latentes sont plausibles. FCI est choisi "
        "plutôt que PC parce qu'il n'exige pas la suffisance causale et renvoie un PAG, c'est-à-dire "
        "une classe de graphes compatibles avec les indépendances conditionnelles détectées [5]. "
        "Cette propriété correspond mieux au risque de promotions, prix ou intentions non mesurés."
    )
    add_body(
        doc,
        "FCI reste un outil exploratoire. Il suppose notamment une structure acyclique compatible avec "
        "les propriétés de Markov et de fidélité, des observations suffisamment indépendantes et des "
        "tests conditionnels valides. La base ne contient qu'une session par utilisateur selon UCI, ce "
        "qui réduit une dépendance évidente entre visites répétées, mais le conditionnement sur "
        "l'existence d'une session peut encore créer une sélection que l'algorithme ne résout pas."
    )

    add_heading(doc, "4.2 Variables, test et contraintes temporelles", 2)
    fci_setup = [
        ["Période", "Feb-Mar; May-June; Jul-Aug-Sep; Oct-Nov-Dec", "Niveau temporel 1"],
        ["Jour spécial", "SpecialDay > 0 : oui/non", "Niveau temporel 1"],
        ["Nouveau visiteur", "VisitorType = New_Visitor : oui/non", "Niveau temporel 1"],
        ["Source de trafic", "Types 1, 2, 3; autres", "Niveau temporel 2"],
        ["Weekend", "Oui/non", "Niveau temporel 2"],
        ["Revenue", "Transaction oui/non", "Niveau temporel 3"],
    ]
    add_table(
        doc,
        number="5",
        title="Variables regroupées pour l'analyse FCI",
        headers=["Variable FCI", "Codage", "Contrainte"],
        rows=fci_setup,
        widths=[2000, 4800, 2560],
        font_size=8.9,
        source="Source : paramétrage de la cellule FCI du notebook final; base complète N = 12 330.",
    )
    add_body(
        doc,
        "Les six variables sont discrétisées afin d'utiliser un test G². Le seuil principal est "
        "α = 0,01, la profondeur maximale des ensembles conditionnants est 2 et la longueur "
        "maximale des chemins explorés est 3. Ces limites réduisent la variance des tableaux de "
        "contingence et le coût du calcul, au prix d'une exploration moins exhaustive. Les niveaux "
        "temporels empêchent les variables tardives de causer les variables précoces; certaines "
        "orientations du PAG proviennent donc de connaissances imposées, et non des seules données."
    )
    add_figure(
        doc,
        filename="03_pag_fci.png",
        number=3,
        title="PAG exploratoire obtenu avec FCI",
        source="Source : FCI sur la base complète, test G², α = 0,01, profondeur 2, chemin maximal 3. Notation : o—o indéterminée; o→ partielle; → imposée.",
        alt="PAG à six variables montrant les orientations complètes, partielles et indéterminées; aucune arête directe Weekend-Revenue.",
        width=6.05,
        crop_top=6500,
        crop_bottom=8000,
    )
    stability = [
        ["0,001", "12", "Absente", "JourSpecial; SourceTrafic"],
        ["0,010", "13", "Absente", "JourSpecial; NouveauVisiteur; SourceTrafic"],
        ["0,050", "13", "Absente", "JourSpecial; NouveauVisiteur; SourceTrafic"],
    ]
    add_table(
        doc,
        number="6",
        title="Stabilité de l'adjacence Weekend-Revenue",
        headers=["α", "Arêtes", "Weekend-Revenue", "Voisins de Weekend"],
        rows=stability,
        widths=[1300, 1300, 2500, 4260],
        numeric_cols=(0, 1),
        font_size=8.9,
        source="Source : répétition de FCI aux seuils 0,1 %, 1 % et 5 %.",
    )

    add_heading(doc, "4.3 Interprétation du PAG", 2)
    add_body(
        doc,
        "Aucune adjacence Weekend-Revenue n'apparaît aux trois seuils. Ce résultat indique qu'une "
        "indépendance conditionnelle compatible avec l'absence d'adjacence a été détectée dans cette "
        "représentation simplifiée. Il ne démontre ni un ATE nul, ni l'absence d'un effet indirect, ni "
        "l'absence de biais. Un effet faible, une discrétisation grossière, une erreur de mesure ou une "
        "violation de fidélité peuvent aussi faire disparaître une arête."
    )
    add_body(
        doc,
        "L'arête partiellement orientée Weekend o→ SourceTrafic met en évidence l'ambiguïté de "
        "TrafficType. Le PAG motive donc l'analyse sans cette variable, mais ne choisit pas lui-même "
        "l'ensemble d'ajustement. Le DAG de domaine demeure l'objet utilisé pour l'identification."
    )

    # 5. Framework and identification
    add_heading(doc, "5. Cadre causal et stratégie d'identification", 1)
    add_heading(doc, "5.1 Pourquoi combiner Pearl et Rubin\u00A0?", 2)
    add_body(
        doc,
        "Le cadre de Rubin fournit les résultats potentiels Y(1) et Y(0), définit l'effet moyen et "
        "clarifie les hypothèses d'échangeabilité, de consistance et de positivité [4]. Le cadre de "
        "Pearl représente les mécanismes par un DAG et formalise l'intervention do(T=t) ainsi que le "
        "critère de porte arrière [3]. La combinaison convient à une exposition et à un résultat "
        "binaires, avec une question portant sur un effet total moyen dans une population définie."
    )

    add_heading(doc, "5.2 Résultats potentiels et estimand", 2)
    add_body(
        doc,
        "Pour la session i, Yᵢ(1) représente la transaction potentielle si la session avait lieu le "
        "weekend, et Yᵢ(0) la transaction potentielle si elle avait lieu en semaine. Un seul résultat "
        "est observé. L'estimand principal est l'ATE dans la population restreinte, ni un ATT, ni un "
        "CATE, ni un effet individuel :"
    )
    add_equation(doc, "τ₀ = E[Y(1) − Y(0) | SpecialDay = 0]")
    add_body(
        doc,
        "Puisque Y est binaire, τ₀ est une différence moyenne de probabilités. Une valeur de 0,015 "
        "correspond à 1,5 point de pourcentage, et non à 1,5 % de croissance relative."
    )

    add_heading(doc, "5.3 Hypothèses", 2)
    assumptions_num = add_numbering_definition(doc, kind="decimal")
    add_list_item(doc, "le résultat observé correspond au résultat potentiel associé à l'exposition effectivement reçue, soit Y = Y(T).", assumptions_num, label="Consistance")
    add_list_item(doc, "après ajustement sur Z dans la population SpecialDay=0, aucune cause commune non mesurée de T et Y ne doit subsister.", assumptions_num, label="Échangeabilité conditionnelle")
    add_list_item(doc, "pour presque toute valeur z présente dans la population cible, 0 < P(T=1 | Z=z, SpecialDay=0) < 1.", assumptions_num, label="Positivité")
    add_list_item(doc, "le résultat potentiel d'une session ne doit pas dépendre du statut Weekend d'une autre session.", assumptions_num, label="Absence d'interférence")
    add_list_item(doc, "les covariables d'ajustement doivent être mesurées sans erreur déterminante et avant l'exposition selon le mécanisme supposé.", assumptions_num, label="Mesure et temporalité")
    add_equation(doc, "(Y(1), Y(0)) ⟂ T | Z, SpecialDay = 0", after=4)
    add_equation(doc, "0 < P(T = 1 | Z, SpecialDay = 0) < 1")
    add_body(
        doc,
        "La consistance est délicate parce que Weekend résume plusieurs versions possibles de "
        "l'exposition : samedi ou dimanche, heure, promotions et expérience du site. L'échangeabilité "
        "est l'hypothèse la plus fragile, car les prix, les produits consultés, les campagnes et "
        "l'intention préalable sont absents. La restriction à SpecialDay=0 rend la positivité plus "
        "plausible, et les diagnostics de score de propension l'évaluent empiriquement."
    )

    add_heading(doc, "5.4 Formule d'identification", 2)
    add_body(
        doc,
        "Si Z suffit à bloquer les chemins de porte arrière et si les hypothèses précédentes sont "
        "valides, l'effet est identifié par standardisation sur la distribution de Z dans la population "
        "restreinte. En posant mₜ(z) = E[Y | T=t, Z=z, SpecialDay=0] :"
    )
    add_equation(doc, "τ₀ = E[m₁(Z) − m₀(Z) | SpecialDay = 0]")
    add_equation(doc, "P(Y=1 | do(T=t), SpecialDay=0)")
    add_equation(doc, "= Σ_z P(Y=1 | T=t, Z=z, SpecialDay=0) P(Z=z | SpecialDay=0)")
    add_body(
        doc,
        "Cette identification est conditionnelle, pas démontrée par les données. Si une cause latente "
        "reliant Weekend et Revenue demeure active après ajustement, aucun ensemble de covariables "
        "pré-exposition entièrement observé dans cette base ne ferme le chemin correspondant. Les "
        "estimations qui suivent quantifient donc l'effet sous l'hypothèse d'absence de confusion "
        "résiduelle; l'E-value évalue ensuite la fragilité "
        "de cette hypothèse sur l'échelle des rapports de risques."
    )

    # 6. Estimation and diagnostics
    add_heading(doc, "6. Estimation et diagnostics", 1)
    add_heading(doc, "6.1 Modèles de nuisance et prédictions croisées", 2)
    add_body(
        doc,
        "Le traitement et le résultat étant binaires, les modèles principaux sont des régressions "
        "logistiques. Elles produisent directement des probabilités estimées et restent interprétables "
        "pour un échantillon de plus de onze mille observations. Les covariables catégorielles sont "
        "encodées en indicatrices; les modalités comptant moins de 20 observations sont regroupées par "
        "l'encodeur afin d'éviter des coefficients estimés sur des cellules presque vides."
    )
    add_body(
        doc,
        "Le score e(z)=P(T=1 | Z=z, SpecialDay=0) et les régressions de résultat m₁(Z), m₀(Z) sont estimés par validation "
        "croisée en cinq plis. La stratification conserve les quatre combinaisons de T et Y. Chaque "
        "observation reçoit des prédictions provenant de modèles qui ne l'ont pas utilisée pour "
        "l'entraînement, ce qui limite le surajustement. Les scores sont bornés numériquement dans "
        "[0,01 ; 0,99] pour éviter une division instable; la distribution observée reste ensuite auditée."
    )

    methods = [
        ["Brute", "Différence directe des taux observés", "Référence descriptive, non causale", "Approximation normale"],
        ["Standardisation", "Moyenne des prédictions sous T=1 et T=0", "Lisible et adaptée à un résultat binaire", "Bootstrap percentile, 500 réplications"],
        ["IPW de Hájek", "Moyennes pondérées normalisées par le score", "Estime le même estimand par repondération de la distribution mesurée de Z", "Fonction d'influence"],
        ["AIPW", "Régression du résultat + score de propension", "Estimateur principal doublement robuste", "Fonction d'influence"],
    ]
    add_table(
        doc,
        number="7",
        title="Estimateurs comparés et calcul de l'incertitude",
        headers=["Méthode", "Principe", "Rôle", "IC à 95 %"],
        rows=methods,
        widths=[1600, 3100, 2860, 1800],
        font_size=8.6,
        source="Source : implémentation du notebook final.",
    )

    add_heading(doc, "6.2 Formules des estimateurs", 2)
    add_body(
        doc,
        "La standardisation compare, pour chaque profil observé, les prédictions du même modèle sous "
        "weekend et sous semaine, puis moyenne les contrastes :"
    )
    add_equation(doc, "τ̂ (standardisation) = (1/n) Σᵢ [m̂₁(Zᵢ) − m̂₀(Zᵢ)]")
    add_body(
        doc,
        "L'IPW de Hájek donne davantage de poids aux observations dont l'exposition était moins "
        "probable compte tenu de Z. La normalisation séparée des deux groupes stabilise les moyennes :"
    )
    add_equation(doc, "μ̂₁ (IPW) = [Σᵢ TᵢYᵢ/êᵢ] / [Σᵢ Tᵢ/êᵢ]", after=3)
    add_equation(doc, "μ̂₀ (IPW) = [Σᵢ (1−Tᵢ)Yᵢ/(1−êᵢ)] / [Σᵢ (1−Tᵢ)/(1−êᵢ)]")
    add_body(
        doc,
        "L'AIPW ajoute aux prédictions de résultat une correction pondérée par les résidus observés. "
        "Avec des prédictions croisées, l'estimateur utilisé est :"
    )
    add_equation(
        doc,
        "τ̂ (AIPW) = (1/n) Σᵢ {m̂₁(Zᵢ) − m̂₀(Zᵢ)\n"
        "+ Tᵢ[Yᵢ−m̂₁(Zᵢ)]/êᵢ − (1−Tᵢ)[Yᵢ−m̂₀(Zᵢ)]/(1−êᵢ)}",
    )
    add_body(
        doc,
        "L'AIPW est doublement robuste : sous les conditions régulières, il peut rester cohérent si "
        "le modèle du traitement ou le modèle du résultat est correctement spécifié. Cette propriété "
        "ne protège pas contre un mauvais DAG, une covariable post-traitement, une violation de "
        "positivité ou un confondeur non mesuré [6]."
    )

    add_heading(doc, "6.3 Incertitude", 2)
    add_body(
        doc,
        "L'intervalle de la différence brute utilise l'approximation normale de deux proportions. "
        "La standardisation utilise les quantiles 2,5 % et 97,5 % d'un bootstrap de 500 réplications; "
        "chaque réplication réajuste une régression logistique, sans refaire le cross-fitting complet. "
        "Les erreurs standards IPW et AIPW reposent sur la dispersion empirique des fonctions "
        "d'influence. Ces intervalles décrivent l'incertitude d'échantillonnage conditionnellement aux "
        "modèles; ils n'intègrent pas l'incertitude sur le DAG ni la confusion cachée.",
        keep_together=True,
    )

    add_heading(doc, "6.4 Recouvrement des scores", 2)
    add_figure(
        doc,
        filename="04_recouvrement_propension.png",
        number=4,
        title="Recouvrement des scores de propension",
        source="Source : prédictions croisées du modèle de traitement dans la population SpecialDay=0.",
        alt="Histogrammes superposés des scores de propension pour les sessions de semaine et de weekend, avec seuils 0,05 et 0,95.",
        width=5.95,
        crop_top=6500,
    )
    diagnostics = [
        ["Score minimal", "0,048"],
        ["Score médian", "0,229"],
        ["Score maximal", "0,919"],
        ["Observations hors [0,05 ; 0,95]", "9"],
        ["Poids stabilisé maximal", "7,274"],
        ["Taille effective - semaine", "8 181,525"],
        ["Taille effective - weekend", "2 411,922"],
    ]
    add_table(
        doc,
        number="8",
        title="Diagnostics du score de propension et des poids",
        headers=["Diagnostic", "Valeur"],
        rows=diagnostics,
        widths=[6500, 2860],
        numeric_cols=(1,),
        font_size=9.4,
        source="Source : calculs réalisés à partir des scores de propension croisés.",
    )
    add_body(
        doc,
        "Les distributions se chevauchent sur l'essentiel du support. Neuf observations seulement se "
        "situent hors de [0,05 ; 0,95], et le poids stabilisé maximal reste inférieur à 7,3. Les "
        "tailles effectives conservent environ 97,7 % de l'effectif de semaine et 89,1 % de l'effectif "
        "de weekend. Ces diagnostics soutiennent la positivité mesurée dans la population restreinte, "
        "sans prouver la positivité pour des combinaisons non observées de covariables."
    )

    add_heading(doc, "6.5 Équilibre des covariables", 2)
    add_figure(
        doc,
        filename="05_equilibre_covariables.png",
        number=5,
        title="Équilibre avant et après pondération IPW",
        source="Source : différences moyennes standardisées des indicatrices de covariables.",
        alt="Graphique des différences moyennes standardisées avant et après IPW pour les 18 indicatrices les plus déséquilibrées.",
        width=4.85,
        crop_top=4500,
    )
    max_before = balance["Avant"].abs().max()
    max_after = balance["Après IPW"].abs().max()
    n_before = int((balance["Avant"].abs() > 0.10).sum())
    n_after = int((balance["Après IPW"].abs() > 0.10).sum())
    add_body(
        doc,
        f"Le SMD absolu maximal passe de {fr(max_before)} à {fr(max_after)}. Le nombre d'indicatrices "
        f"dont |SMD| dépasse 0,10 passe de {n_before} à {n_after}. L'IPW obtient ainsi un bon équilibre "
        "marginal des indicatrices mesurées selon le seuil de 0,10. Ce contrôle ne porte pas sur les promotions, les prix ou l'intention "
        "préalable absents de la base."
    )

    # 7. Results
    add_heading(doc, "7. Résultats et analyses de robustesse", 1)
    add_heading(doc, "7.1 Estimations principales", 2)
    main_rows = []
    for _, row in principal.iterrows():
        main_rows.append([
            row["Méthode"],
            fr(row["Estimation (points)"]),
            fr(row["IC 95 % bas"]),
            fr(row["IC 95 % haut"]),
        ])
    add_table(
        doc,
        number="9",
        title="Estimation du contraste semaine-weekend",
        headers=["Méthode", "Estimation (points)", "IC 95 % bas", "IC 95 % haut"],
        rows=main_rows,
        widths=[3000, 2240, 2060, 2060],
        numeric_cols=(1, 2, 3),
        font_size=9.1,
        source="Source : résultats produits par le notebook final.",
    )
    add_figure(
        doc,
        filename="06_estimations_principales.png",
        number=6,
        title="Association brute et estimations ajustées",
        source="Source : notebook final; les segments représentent les intervalles à 95 %.",
        alt="Graphique en points avec intervalles pour l'association brute, la standardisation, l'IPW et l'AIPW.",
        width=5.95,
        crop_top=6500,
    )
    add_body(
        doc,
        "Les trois méthodes ajustées donnent des points estimés proches, de +1,384 à +1,506 point. "
        "L'AIPW estime une probabilité moyenne de 17,67 % sous weekend et de 16,16 % sous semaine, "
        "soit +1,506 point. Son intervalle à 95 % [-0,176 ; 3,188] contient zéro. La cohérence des "
        "points estimés est rassurante sur la stabilité numérique, mais elle ne permet pas d'écarter "
        "l'absence d'effet moyen."
    )

    add_heading(doc, "7.2 Signification pratique", 2)
    add_body(
        doc,
        "Sous toutes les hypothèses causales, +1,51 point correspondrait à environ 15 transactions "
        "supplémentaires pour 1 000 sessions comparables déplacées vers le weekend. L'intervalle "
        "principal correspond approximativement à -2 à +32 transactions pour 1 000 sessions. Le "
        "rapport de risques estimé est 1,093, soit environ 9,3 % de plus relativement au risque "
        "potentiel sous semaine. Ces conversions concernent un nombre fixe de sessions existantes; "
        "elles ne prédisent pas un gain de trafic ou de chiffre d'affaires."
    )

    add_heading(doc, "7.3 Spécifications alternatives", 2)
    robust_rows = []
    for _, row in robustness.iterrows():
        robust_rows.append([
            row["Scénario"],
            f"{int(row['N']):,}".replace(",", " "),
            fr(row["Estimation (points)"]),
            f"[{fr(row['IC 95 % bas'])} ; {fr(row['IC 95 % haut'])}]",
        ])
    add_table(
        doc,
        number="10",
        title="Analyses de sensibilité de l'AIPW",
        headers=["Scénario", "N", "Estimation", "IC à 95 %"],
        rows=robust_rows,
        widths=[4600, 1100, 1400, 2260],
        numeric_cols=(1, 2, 3),
        font_size=8.6,
        source="Source : analyses de robustesse exécutées dans le notebook final.",
    )
    add_figure(
        doc,
        filename="07_robustesse_aipw.png",
        number=7,
        title="Sensibilité de l'estimation AIPW",
        source="Source : six spécifications de population, d'ajustement et de modèles de nuisance.",
        alt="Graphique en points des six analyses AIPW avec leurs intervalles à 95 %.",
        width=6.05,
        crop_top=5000,
    )
    add_body(
        doc,
        "Les estimations demeurent positives, de +1,193 à +2,233 points. Retirer TrafficType produit "
        "l'estimation la plus élevée, +2,233 [0,535 ; 3,931], ce qui montre que l'estimation est "
        "sensible à la manière dont son statut causal est défini. La forêt aléatoire donne +1,736 "
        "[0,123 ; 3,348]. Ces deux intervalles excluent "
        "zéro, mais il s'agit de spécifications alternatives : elles n'annulent ni l'incertitude de "
        "l'analyse principale ni la variabilité entre choix de modélisation."
    )
    add_body(
        doc,
        "La population complète, le retrait des doublons et la restriction au support commun conduisent "
        "à la même conclusion prudente, avec des intervalles qui contiennent zéro. Aucun contrôle négatif "
        "ou placebo crédible n'est disponible dans cette base; l'absence d'une telle vérification est "
        "préférable à l'invention d'un contrôle dont les conditions causales seraient contestables."
    )

    add_heading(doc, "7.4 Sensibilité à la confusion non mesurée", 2)
    add_body(
        doc,
        "Les moyennes AIPW donnent un rapport de risques de 1,093. L'E-value du point estimé vaut "
        "1,412 : sous la logique de cette mesure, une cause non observée ayant des associations d'au "
        "moins 1,41 avec l'exposition et avec le résultat, sur l'échelle des rapports de risques et "
        "au-delà des covariables mesurées, pourrait suffire à ramener le point vers le nul [7]. "
        "L'E-value de la borne la plus proche du nul vaut 1,000 parce que l'intervalle "
        "principal contient déjà zéro. Cette valeur indique une robustesse limitée; elle ne mesure pas "
        "la probabilité qu'un confondeur existe et ne corrige pas l'estimation."
    )

    # 8. Discussion
    add_heading(doc, "8. Discussion et limites", 1)
    add_heading(doc, "8.1 Ce que les résultats permettent d'affirmer", 2)
    add_body(
        doc,
        "Dans la population SpecialDay=0, les sessions de weekend présentent un taux de transaction "
        "plus élevé. Après ajustement sur les covariables pré-session mesurées, le contraste demeure "
        "positif et plus faible que l'association brute. Le bon recouvrement, la réduction des SMD et "
        "la proximité des estimateurs ne suggèrent pas que le résultat soit principalement entraîné par "
        "un petit nombre de poids extrêmes ou par un déséquilibre résiduel évident des variables mesurées."
    )
    add_body(
        doc,
        "La conclusion causale reste néanmoins faible. L'intervalle AIPW principal contient zéro et "
        "l'estimation varie selon le statut de TrafficType et le modèle de nuisance. La formulation la "
        "plus défendable est donc la suivante : les données sont compatibles avec un faible contraste "
        "positif du weekend sur la probabilité de transaction, mais elles ne permettent pas d'établir un effet "
        "causal moyen différent de zéro."
    )

    add_heading(doc, "8.2 Menaces à la validité", 2)
    limits_num = add_numbering_definition(doc, kind="decimal")
    add_list_item(doc, "Weekend résume plusieurs versions du traitement. Une session du samedi matin et une session du dimanche soir ne constituent pas nécessairement la même intervention.", limits_num, label="Intervention imparfaitement définie")
    add_list_item(doc, "prix, promotions, produits, campagnes, heure et intention d'achat préalable sont absents; ils peuvent expliquer une partie ou la totalité du contraste.", limits_num, label="Confusion non mesurée")
    add_list_item(doc, "la base ne contient que des visites observées. Une politique de weekend peut aussi modifier la probabilité qu'une session existe, effet non estimé ici.", limits_num, label="Sélection sur les sessions")
    add_list_item(doc, "TrafficType peut être un confondeur, un médiateur ou un indicateur imparfait d'une campagne. La sensibilité à son retrait confirme cette incertitude.", limits_num, label="Ordre causal ambigu")
    add_list_item(doc, "les catégories de région, de système, de navigateur et de trafic sont anonymes, ce qui limite la validation externe des mécanismes.", limits_num, label="Mesure et connaissance du domaine")
    add_list_item(doc, "FCI dépend de la discrétisation, du test G², de la profondeur, des contraintes temporelles et des hypothèses de Markov et de fidélité.", limits_num, label="Découverte causale")
    add_list_item(doc, "le bootstrap et les fonctions d'influence quantifient l'incertitude d'échantillonnage, pas l'incertitude sur le DAG ou sur les variables absentes.", limits_num, label="Inférence statistique")
    add_list_item(doc, "le résultat s'applique au site et à l'année observés, pour des sessions hors jours spéciaux. Il ne se transpose pas automatiquement à d'autres plateformes ou périodes.", limits_num, label="Validité externe")

    add_heading(doc, "8.3 Prochaine étude recommandée", 2)
    add_body(
        doc,
        "Une étude plus crédible collecterait l'horodatage, les prix, les promotions, le produit, la "
        "campagne, l'appareil et des indicateurs d'intention mesurés avant la session. Une expérience "
        "pourrait randomiser un levier réellement manipulable lié au calendrier, par exemple l'heure "
        "d'envoi d'une campagne, tout en mesurant à la fois la probabilité de visite et la conversion. "
        "À défaut, un plan longitudinal avec effets de calendrier plus fins et contrôles négatifs "
        "préspécifiés permettrait de tester certaines menaces sans prétendre les éliminer."
    )

    # 9. Reproducibility
    add_heading(doc, "9. Reproductibilité", 1)
    add_body(
        doc,
        "Le notebook final est exécuté de la première à la dernière cellule. Il vérifie l'empreinte du "
        "CSV avant toute analyse, fixe la graine 20260715, produit les prédictions croisées, exécute "
        "les 500 réplications du bootstrap et régénère les sept figures ainsi que les quatre tableaux "
        "CSV. Les versions logicielles validées sont fixées dans requirements.txt."
    )
    repro_bullets = add_numbering_definition(doc, kind="bullet")
    add_list_item(doc, "Notebook exécuté : 01_notebooks/livrable/projet_final_inference_causale_online_shoppers.ipynb", repro_bullets)
    add_list_item(doc, "Données : 02_data/raw/online_shoppers_intention.csv", repro_bullets)
    add_list_item(doc, "Figures et tableaux : 04_outputs/figures/ et 04_outputs/tableaux/", repro_bullets)
    add_list_item(doc, "Export HTML : 04_outputs/notebook_html/projet_final_inference_causale_online_shoppers.html", repro_bullets)
    add_list_item(doc, "Dépendances : requirements.txt; environnement validé sous Python 3.14.3.", repro_bullets)
    add_body(
        doc,
        "Depuis la racine extraite, l'utilisateur installe les dépendances, puis lance l'exécution "
        "non interactive indiquée dans le README. Les chemins sont relatifs afin que le notebook "
        "retrouve les données et réécrive ses sorties dans la structure livrée."
    )

    # 10. Conclusion
    add_heading(doc, "10. Conclusion", 1)
    add_body(
        doc,
        "Le projet a formulé une question causale précise, défini une exposition, un résultat, une "
        "population et un estimand, puis relié un DAG de domaine, une exploration FCI, une stratégie "
        "de porte arrière et quatre estimations. L'association brute est de +2,03 points dans la "
        "population retenue. L'AIPW donne +1,51 point, avec un IC à 95 % de [-0,18 ; 3,19]."
    )
    add_body(
        doc,
        "Les diagnostics sont favorables pour les covariables mesurées dans la population restreinte, "
        "mais ils ne résolvent pas la confusion non observée "
        "et l'intervalle principal contient zéro. Il n'est donc pas justifié d'affirmer que le weekend "
        "cause une hausse des transactions. Le résultat utile est plus modeste : sous les hypothèses explicitées, "
        "un petit contraste positif est plausible, et l'analyse montre précisément quelles données et "
        "quel plan d'étude seraient nécessaires pour transformer cette hypothèse en conclusion plus solide.",
        keep_together=True,
    )

    # References
    add_heading(doc, "Références", 1)
    refs_num = add_numbering_definition(doc, kind="decimal")
    references = [
        "Sakar, C. et Kastro, Y. (2018). Online Shoppers Purchasing Intention Dataset. UCI Machine Learning Repository. DOI : 10.24432/C5F88Q.",
        "Sakar, C. O., Polat, S. O., Katircioglu, M. et Kastro, Y. (2019). Real-time prediction of online shoppers' purchasing intention using multilayer perceptron and LSTM recurrent neural networks. Neural Computing and Applications, 31, 6893-6908.",
        "Pearl, J. (2009). Causality: Models, Reasoning, and Inference (2e éd.). Cambridge University Press.",
        "Rosenbaum, P. R. et Rubin, D. B. (1983). The central role of the propensity score in observational studies for causal effects. Biometrika, 70(1), 41-55.",
        "Spirtes, P., Glymour, C. et Scheines, R. (2000). Causation, Prediction, and Search (2e éd.). MIT Press.",
        "Bang, H. et Robins, J. M. (2005). Doubly robust estimation in missing data and causal inference models. Biometrics, 61(4), 962-973.",
        "VanderWeele, T. J. et Ding, P. (2017). Sensitivity analysis in observational research: introducing the E-value. Annals of Internal Medicine, 167(4), 268-274.",
    ]
    for ref_text in references:
        add_list_item(doc, ref_text, refs_num)

    # Appendix
    add_heading(doc, "Annexe A. Arêtes du PAG FCI", 1, page_break=True)
    add_body(
        doc,
        "Le tableau A1 reproduit l'ensemble des 13 arêtes du PAG au seuil principal α = 0,01. "
        "Le symbole o—o indique une orientation indéterminée; o→ une orientation partielle; → une "
        "direction compatible avec les contraintes et les indépendances détectées."
    )
    pag_rows = [
        [value.replace("o->", "o→").replace("-->", "→").replace("o-o", "o—o")]
        for value in pag.iloc[:, 0].astype(str).tolist()
    ]
    add_table(
        doc,
        number="A1",
        title="Liste complète des arêtes du PAG au seuil α = 0,01",
        headers=["Arête du PAG"],
        rows=pag_rows,
        widths=[9360],
        font_size=9.3,
        source="Source : sortie FCI exportée par le notebook final.",
        keep_whole=False,
    )

    # Remove the final empty paragraph only if Word inserted one after the table source.
    doc.save(OUTPUT_DOCX)
    harden_ooxml_fonts(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    if len(sys.argv) == 3 and sys.argv[1] == "--harden":
        harden_ooxml_fonts(Path(sys.argv[2]).resolve())
    elif len(sys.argv) == 1:
        build_report()
    else:
        raise SystemExit("Usage: build_report_final.py [--harden PATH]")
